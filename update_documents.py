#!/usr/bin/env python3
"""
╔══════════════════════════════════════════════════════════════════════════════╗
║   TERRASOCIAL — Script Maître de Mise à Jour des Documents (Février 2026)  ║
║   MANO VERDE INC SA                                                         ║
╠══════════════════════════════════════════════════════════════════════════════╣
║   Changements propagés :                                                    ║
║   • Suppression de l'acompte 10% du prix du lot                             ║
║   • Remplacement par frais d'ouverture de dossier : 10 000 FCFA (forfait)  ║
║   • Ajout du mode versement journalier : dès 1 500 FCFA/jour               ║
║   • Slogan : « 2 bières par jour pour votre terrain titré ! »              ║
╚══════════════════════════════════════════════════════════════════════════════╝

Usage :
    python3 update_documents.py [--target /chemin/vers/PROJET_FONCIER_SOCIAL]

Si --target n'est pas spécifié, le script utilise le chemin par défaut :
    ~/Documents/Mano_Verde_SA/MonBot/documents/PROJET_FONCIER_SOCIAL
"""

import os
import sys
import shutil
import argparse
import subprocess
from pathlib import Path
from datetime import datetime

# ── Configuration ─────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).parent.resolve()   # dossier Code_source
DEFAULT_TARGET = Path.home() / 'Documents' / 'Mano_Verde_SA' / 'MonBot' / 'documents' / 'PROJET_FONCIER_SOCIAL'
GENERATE_JS  = SCRIPT_DIR / 'generate_docs.js'
GENERATE_PDF = SCRIPT_DIR / 'generate_prospectus.py'

# Documents générés par ce script (dans SCRIPT_DIR)
GENERATED_DOCS = [
    'CGV_TERRASOCIAL_Fev2026.docx',
    'Politique_Paiement_TERRASOCIAL_Fev2026.docx',
    'Contrat_Reservation_TERRASOCIAL_Fev2026.docx',
    'Note_MiseAJour_NouveauModele_Fev2026.docx',
    'Prospectus_TERRASOCIAL_Fev2026.pdf',
]

# Substitutions de texte à appliquer dans les anciens documents Word
TEXT_REPLACEMENTS = [
    # (ancien texte, nouveau texte, description)
    ('acompte de 10%',       'frais de dossier de 10 000 FCFA',          'Acompte → frais dossier'),
    ('acompte 10%',          'frais de dossier 10 000 FCFA',              'Acompte → frais dossier'),
    ('Acompte : 10%',        'Frais de dossier : 10 000 FCFA',            'Ligne acompte'),
    ('Acompte (10%)',        'Frais de dossier (10 000 FCFA)',             'Label acompte'),
    ('acompte obligatoire',  "frais d'ouverture de dossier",              'Terme acompte'),
    ('10% du prix',          'frais de dossier de 10 000 FCFA',           '10% ref'),
    ('versement de 10%',     'frais de dossier de 10 000 FCFA',           'Versement 10%'),
    ('mensualités uniquement', 'mensualités ou versements journaliers',   'Mode versement'),
    ('mensualite uniquement', 'mensualite ou versements journaliers',     'Mode versement'),
    ('paiement mensuel',     'paiement mensuel ou journalier',            'Mode paiement'),
]

# Termes à signaler (présents dans les docs mais pas remplacés automatiquement)
FLAG_TERMS = ['acompte', '10 %', '10%', 'Acompte', 'mensualité uniquement']

# ── Utilitaires ───────────────────────────────────────────────────────────────
RESET  = '\033[0m'
GREEN  = '\033[32m'
ORANGE = '\033[33m'
RED    = '\033[31m'
BOLD   = '\033[1m'
CYAN   = '\033[36m'

def ok(msg):   print(f'{GREEN}  ✅ {msg}{RESET}')
def warn(msg): print(f'{ORANGE}  ⚠️  {msg}{RESET}')
def err(msg):  print(f'{RED}  ❌ {msg}{RESET}')
def info(msg): print(f'{CYAN}  ℹ️  {msg}{RESET}')
def head(msg): print(f'\n{BOLD}{GREEN}══ {msg} ══{RESET}\n')


def ensure_target(target: Path) -> bool:
    """Crée le dossier cible si nécessaire."""
    if not target.exists():
        try:
            target.mkdir(parents=True, exist_ok=True)
            ok(f'Dossier créé : {target}')
        except Exception as e:
            err(f'Impossible de créer le dossier cible : {e}')
            return False
    else:
        info(f'Dossier cible : {target}')
    return True


def regenerate_documents():
    """Relance la génération des documents Word et PDF."""
    head('GÉNÉRATION DES DOCUMENTS')

    # Word documents (Node.js)
    if GENERATE_JS.exists():
        info('Génération des documents Word...')
        result = subprocess.run(
            ['node', str(GENERATE_JS), str(SCRIPT_DIR)],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            for line in result.stdout.strip().splitlines():
                ok(line.strip())
        else:
            warn(f'Avertissement génération Word : {result.stderr[:200]}')
    else:
        warn(f'Script JS non trouvé : {GENERATE_JS} — utilisation des fichiers existants.')

    # Prospectus PDF (Python)
    if GENERATE_PDF.exists():
        info('Génération du Prospectus PDF...')
        result = subprocess.run(
            [sys.executable, str(GENERATE_PDF), str(SCRIPT_DIR)],
            capture_output=True, text=True
        )
        if result.returncode == 0:
            ok(result.stdout.strip())
        else:
            warn(f'Avertissement génération PDF : {result.stderr[:200]}')
    else:
        warn(f'Script PDF non trouvé : {GENERATE_PDF} — utilisation des fichiers existants.')


def copy_to_target(target: Path) -> list:
    """Copie les documents générés vers le dossier cible."""
    head('COPIE DES DOCUMENTS VERS LE DOSSIER PROJET')
    copied = []
    for doc_name in GENERATED_DOCS:
        src = SCRIPT_DIR / doc_name
        dst = target / doc_name
        if src.exists():
            try:
                shutil.copy2(src, dst)
                size_kb = round(src.stat().st_size / 1024)
                ok(f'{doc_name} ({size_kb} Ko) → {target.name}/')
                copied.append(doc_name)
            except Exception as e:
                err(f'{doc_name} : {e}')
        else:
            warn(f'{doc_name} non trouvé dans {SCRIPT_DIR}')
    return copied


def update_existing_word_docs(target: Path) -> dict:
    """
    Cherche et met à jour les anciens documents Word dans le dossier cible.
    Remplace les termes liés à l'acompte 10% par le nouveau modèle.
    Retourne un dict {fichier: [modifications]}.
    """
    head('MISE À JOUR DES DOCUMENTS WORD EXISTANTS')

    try:
        from docx import Document as DocxDocument
    except ImportError:
        warn("python-docx non installé. Tentative d'installation...")
        subprocess.run([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'], check=True)
        from docx import Document as DocxDocument

    results = {}
    docx_files = list(target.rglob('*.docx'))

    if not docx_files:
        info('Aucun fichier .docx trouvé dans le dossier cible.')
        return results

    for docx_path in docx_files:
        # Ne pas modifier les fichiers qu'on vient de générer
        if docx_path.name in GENERATED_DOCS:
            continue

        try:
            doc = DocxDocument(str(docx_path))
            modifications = []
            flagged = []
            changed = False

            # Parcourir tous les paragraphes
            for para in doc.paragraphs:
                for old, new, desc in TEXT_REPLACEMENTS:
                    if old.lower() in para.text.lower():
                        for run in para.runs:
                            if old.lower() in run.text.lower():
                                run.text = run.text.replace(old, new)
                                run.text = run.text.replace(old.lower(), new.lower())
                                changed = True
                                if desc not in modifications:
                                    modifications.append(desc)

                # Termes à signaler
                for term in FLAG_TERMS:
                    if term.lower() in para.text.lower() and term not in [m for m in modifications]:
                        if term not in flagged:
                            flagged.append(f'Terme "{term}" trouvé (vérification manuelle conseillée)')

            # Parcourir les tableaux
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for old, new, desc in TEXT_REPLACEMENTS:
                                if old.lower() in para.text.lower():
                                    for run in para.runs:
                                        if old.lower() in run.text.lower():
                                            run.text = run.text.replace(old, new)
                                            run.text = run.text.replace(old.lower(), new.lower())
                                            changed = True
                                            if desc not in modifications:
                                                modifications.append(desc)

            if changed:
                # Sauvegarde avec suffixe _MisAJour
                stem = docx_path.stem
                if not stem.endswith('_v2') and 'Fev2026' not in stem:
                    new_path = docx_path.with_name(f'{stem}_MisAJour_Fev2026.docx')
                else:
                    new_path = docx_path
                doc.save(str(new_path))
                ok(f'{docx_path.name} → {new_path.name}')
                for mod in modifications:
                    info(f'    • {mod}')
                results[docx_path.name] = modifications
            else:
                info(f'{docx_path.name} : aucune modification automatique nécessaire')
                if flagged:
                    for flag in flagged:
                        warn(f'    {flag}')

        except Exception as e:
            err(f'Erreur sur {docx_path.name} : {e}')

    return results


def generate_rapport(target: Path, copied: list, updated: dict):
    """Génère un rapport de mise à jour dans le dossier cible."""
    head('RAPPORT DE MISE À JOUR')

    now = datetime.now().strftime('%d/%m/%Y à %H:%M')
    rapport_path = target / f'RAPPORT_MiseAJour_TERRASOCIAL_{datetime.now().strftime("%Y%m%d")}.txt'

    lines = [
        '=' * 70,
        'RAPPORT DE MISE À JOUR — TERRASOCIAL / MANO VERDE INC SA',
        f'Date : {now}',
        '=' * 70,
        '',
        'CHANGEMENTS PRINCIPAUX APPLIQUÉS',
        '-' * 40,
        "• Suppression de l'acompte 10% du prix du lot",
        '• Nouveau : Frais d\'ouverture de dossier = 10 000 FCFA (forfait unique)',
        '• Nouveau : Mode versement journalier dès 1 500 FCFA/jour',
        '• Slogan : "2 bières par jour pour votre terrain titré !"',
        '',
        'DOCUMENTS COPIÉS/GÉNÉRÉS',
        '-' * 40,
    ]
    for doc in copied:
        lines.append(f'  ✅ {doc}')
    if not copied:
        lines.append('  Aucun document copié.')

    lines += ['', 'DOCUMENTS MIS À JOUR (anciens fichiers)', '-' * 40]
    if updated:
        for fname, mods in updated.items():
            lines.append(f'  ✅ {fname}')
            for mod in mods:
                lines.append(f'      • {mod}')
    else:
        lines.append('  Aucun fichier ancien mis à jour (ou aucun trouvé).')

    lines += [
        '',
        'RÉSUMÉ DU NOUVEAU MODÈLE',
        '-' * 40,
        '',
        '  ANCIEN MODÈLE :',
        '    - Acompte 10% du prix du lot (ex: 50 000 FCFA pour lot à 500 000 FCFA)',
        '    - Mensualités uniquement',
        '',
        '  NOUVEAU MODÈLE :',
        '    - Frais de dossier : 10 000 FCFA (fixe, quel que soit le lot)',
        '    - Prix du lot payable en mensualités OU versements journaliers',
        '    - Minimum journalier : 1 500 FCFA/jour',
        '    - Paiements anticipés autorisés',
        '',
        'DOCUMENTS DE RÉFÉRENCE',
        '-' * 40,
        '  • CGV_TERRASOCIAL_Fev2026.docx',
        '  • Politique_Paiement_TERRASOCIAL_Fev2026.docx',
        '  • Contrat_Reservation_TERRASOCIAL_Fev2026.docx',
        '  • Note_MiseAJour_NouveauModele_Fev2026.docx',
        '  • Prospectus_TERRASOCIAL_Fev2026.pdf',
        '',
        '=' * 70,
        'Produit automatiquement par update_documents.py — TERRASOCIAL',
        '=' * 70,
    ]

    rapport_path.write_text('\n'.join(lines), encoding='utf-8')
    ok(f'Rapport généré : {rapport_path.name}')
    return rapport_path


def main():
    parser = argparse.ArgumentParser(
        description='TERRASOCIAL — Mise à jour des documents du projet'
    )
    parser.add_argument(
        '--target', '-t',
        default=str(DEFAULT_TARGET),
        help=f'Chemin du dossier PROJET_FONCIER_SOCIAL (défaut: {DEFAULT_TARGET})'
    )
    parser.add_argument(
        '--no-regen', action='store_true',
        help='Ne pas régénérer les documents (utiliser les fichiers existants)'
    )
    parser.add_argument(
        '--no-update-existing', action='store_true',
        help='Ne pas mettre à jour les anciens documents Word existants'
    )
    args = parser.parse_args()

    target = Path(args.target)

    print(f'\n{BOLD}🌿 TERRASOCIAL — Script Maître de Mise à Jour{RESET}')
    print(f'{BOLD}   MANO VERDE INC SA — Février 2026{RESET}')
    print(f'   Source  : {SCRIPT_DIR}')
    print(f'   Cible   : {target}')
    print()

    # 1. Créer dossier cible
    if not ensure_target(target):
        sys.exit(1)

    # 2. Régénérer les documents
    if not args.no_regen:
        regenerate_documents()
    else:
        info('Régénération ignorée (--no-regen)')

    # 3. Copier vers le dossier cible
    copied = copy_to_target(target)

    # 4. Mettre à jour les anciens documents
    updated = {}
    if not args.no_update_existing:
        updated = update_existing_word_docs(target)
    else:
        info('Mise à jour des anciens docs ignorée (--no-update-existing)')

    # 5. Rapport
    rapport = generate_rapport(target, copied, updated)

    print(f'\n{BOLD}{GREEN}═══ TERMINÉ ═══{RESET}')
    print(f'  {len(copied)} document(s) copié(s)')
    print(f'  {len(updated)} document(s) ancien(s) mis à jour')
    print(f'  Rapport : {rapport}')
    print(f'\n  📁 Ouvrez : {target}\n')


if __name__ == '__main__':
    main()
